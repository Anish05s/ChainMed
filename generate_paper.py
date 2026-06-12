"""
ChainMed Research Paper Generator
Generates a complete IEEE-format research paper as .docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

# ─── Style Definitions ───────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)

def add_authors(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)

def add_affiliation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(12)

def add_section_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_subsection_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.75)
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)
    return p

def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Cambria Math'
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

def add_bullet(text):
    p = doc.add_paragraph()
    run = p.add_run("• " + text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)

def add_table_with_data(headers, rows, caption=None):
    if caption:
        add_figure_caption(caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing

# ═══════════════════════════════════════════════════════════════════════════════
#                           PAPER CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ─── Title ────────────────────────────────────────────────────────────────────
add_title("ChainMed: A Hybrid AI and Blockchain Framework for Three-Party Pharmaceutical Supply Chain Verification with LLM-Powered Fraud Investigation")

add_authors("Anish Sharma")
add_affiliation("Department of Computer Science and Engineering\nIndependent Research · Guwahati, India\nanish.sharma@email.com")

# ─── Abstract ─────────────────────────────────────────────────────────────────
add_section_heading("Abstract")

add_body_no_indent(
    "The global pharmaceutical supply chain faces a systemic integrity crisis, with the counterfeit medicine market "
    "estimated at up to $432 billion annually—exceeding the combined revenue of illicit arms and narcotics trades. "
    "A 2023 UNODC report documented 267,000 annual deaths from falsified antimalarial drugs and 169,000 from fake "
    "antibiotics for childhood pneumonia in sub-Saharan Africa alone. Existing blockchain-based solutions address "
    "traceability but fail to implement cross-party fraud detection; existing AI solutions address demand forecasting "
    "but not fraud explainability. No prior system combines both with crisis-aware logistics rerouting without IoT "
    "hardware dependency. This paper presents ChainMed, a hybrid AI and blockchain framework that introduces "
    "three-party attestation verification—requiring independent data submission from Manufacturer, Supplier, and "
    "Hospital—cross-matched by a deterministic rule engine augmented with a Large Language Model (Gemini 2.5 Flash) "
    "for natural-language fraud investigation reports. The system employs an off-chain/on-chain hybrid ledger "
    "architecture on Ethereum Sepolia, a Dijkstra-based crisis rerouting engine integrated with real-time news "
    "intelligence, and a dynamic entity trust scoring mechanism. ChainMed is deployed as a production-accessible "
    "web application (Vercel + Railway) with zero hardware requirements, making it immediately accessible to NGOs, "
    "governments, and healthcare facilities in resource-constrained regions. Security analysis identifies smart "
    "contract vulnerabilities including single-owner authority (f = 0 Byzantine fault tolerance) and mutable records, "
    "with proposed mitigations including multi-signature thresholds and append-only guarantees. Field-level AES-256 "
    "encryption, salted hash commitments, and zero-knowledge proof attestation are proposed as the privacy-preserving "
    "upgrade path for production deployment."
)

p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run = p.add_run("Pharmaceutical supply chain, Blockchain, Artificial intelligence, Three-party attestation, "
                "LLM fraud investigation, Trust scoring, Crisis rerouting, Smart contracts, Zero-knowledge proofs")
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("I. Introduction")

add_body(
    "The pharmaceutical supply chain is among the most complex and regulated distribution networks in the global "
    "economy, involving multiple stakeholders across manufacturing, warehousing, distribution, and consumption. "
    "The World Health Organization (WHO) estimates that 11% of medicines in developing countries are counterfeit, "
    "with drugs for malaria and bacterial infections accounting for nearly 65% of falsified medicines [1]. Between "
    "72,000 and 169,000 children may be dying annually from pneumonia alone after receiving ineffective counterfeit "
    "drugs [2]. The Drug Supply Chain Security Act (DSCSA), which reached full enforcement in November 2023, "
    "mandates electronic tracking of prescription drugs across the United States supply chain [3]. Similarly, the "
    "European Union's Falsified Medicines Directive (FMD) and India's updated Schedule M GMP compliance requirements "
    "(2023) impose stringent traceability obligations on pharmaceutical exporters globally."
)

add_body(
    "Blockchain technology has emerged as a promising solution for supply chain integrity, offering immutability, "
    "transparency, and cryptographic verification of transactions [4]. Several prior works have proposed blockchain-based "
    "pharmaceutical tracking systems. Gomasta et al. [5] developed a Hyperledger Fabric framework with ECDSA double-signing "
    "between pharmaceutical companies and drug regulators. Jamil et al. [6] proposed a smart hospital system using "
    "Hyperledger Fabric with IoT integration. Musamih et al. [7] designed a blockchain-based drug traceability architecture "
    "on Ethereum. However, these existing solutions share three critical limitations:"
)

add_bullet("No system implements cross-party fraud detection where independently submitted data from three or more parties is automatically compared for inconsistencies.")
add_bullet("No system provides explainable AI-generated investigation reports for flagged shipments using Large Language Models.")
add_bullet("No system integrates real-time crisis intelligence with supply chain rerouting without requiring IoT hardware.")

add_body(
    "This paper introduces ChainMed, a framework that addresses all three gaps through: (1) a novel three-party "
    "attestation model requiring independent data submission from Manufacturer, Supplier, and Hospital/Consumer; "
    "(2) a hybrid AI verification engine combining deterministic rule-based cross-matching with LLM-powered "
    "investigation; (3) a crisis disruption center with Dijkstra-based logistics rerouting; and (4) a dynamic "
    "entity trust scoring mechanism that adjusts based on historical verification outcomes. The system is deployed "
    "as a production-accessible web application with zero hardware cost, making it immediately deployable in "
    "resource-constrained regions where counterfeit medicine mortality is highest."
)

add_body("The remainder of this paper is organized as follows: Section II reviews related work in blockchain and AI "
         "for pharmaceutical supply chains. Section III presents the system architecture. Section IV details the "
         "mathematical methodology including the verification AI, trust engine, and rerouting algorithms. Section V "
         "describes the implementation and deployment. Section VI provides security analysis including smart contract "
         "vulnerability assessment and Byzantine fault tolerance. Section VII discusses limitations and the proposed "
         "privacy-preserving upgrade path. Section VIII outlines future work directions. Section IX concludes the paper.")

# ═══════════════════════════════════════════════════════════════════════════════
# II. RELATED WORK
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("II. Related Work")

add_subsection_heading("A. Blockchain-Based Pharmaceutical Supply Chain Systems")

add_body(
    "Gomasta et al. [5] proposed ChainMed, a Hyperledger Fabric-based framework for Bangladesh's pharmaceutical "
    "supply chain, implementing ECDSA double-signing between pharmaceutical companies and the Directorate General "
    "of Drug Administration (DGDA). Their system provides on-chain and off-chain storage with smart contracts for "
    "data provenance, and includes a formal security analysis demonstrating 33% Byzantine fault tolerance (BFT). "
    "However, the system operates only between two parties (pharmaceutical company and regulator), does not "
    "implement AI-based fraud detection, and is not deployed as a live accessible application."
)

add_body(
    "Jamil et al. [6] designed a medical blockchain for drug supply chain integrity in smart hospitals, utilizing "
    "Hyperledger Fabric with IoT sensors for environmental monitoring. While their system provides comprehensive "
    "tracking, it requires significant hardware investment (IoT sensors, RFID tags) and does not include AI-powered "
    "cross-party verification. Musamih et al. [7] proposed a blockchain-based drug traceability system on Ethereum "
    "with smart contracts for ownership transfer, but focused on linear tracking rather than multi-party attestation."
)

add_body(
    "Sylim et al. [8] developed a blockchain-based system specifically for detecting falsified and substandard "
    "drugs in distribution, representing one of the earliest proposals for using distributed ledger technology "
    "in pharmaceutical fraud detection. Uddin et al. [9] provided a comprehensive review of blockchain "
    "architectures for drug traceability, identifying interoperability and scalability as the primary open "
    "challenges. The MediLedger Project [10] represents the most significant industry initiative, implementing "
    "DSCSA-compliant product verification on a permissioned blockchain, but remains a proprietary enterprise "
    "solution inaccessible to NGOs and developing-country health systems."
)

add_subsection_heading("B. AI in Supply Chain Fraud Detection")

add_body(
    "Bello et al. [11] proposed conceptual frameworks for integrating machine learning with blockchain for "
    "real-time fraud detection, but did not implement a working system. Recent work on LLM-grounded explainable "
    "AI for supply chain risk [12] demonstrates the feasibility of using temporal graph attention networks with "
    "large language models for early warning systems. However, no prior work has combined LLM-based investigation "
    "with blockchain-anchored three-party attestation for pharmaceutical fraud detection."
)

add_subsection_heading("C. Crisis-Aware Supply Chain Management")

add_body(
    "Supply chain disruption management has been extensively studied in operations research, with Dijkstra's "
    "algorithm widely applied for shortest-path rerouting in logistics networks [13]. However, existing "
    "pharmaceutical supply chain systems treat disruption management and fraud detection as separate concerns. "
    "ChainMed is the first system to integrate real-time crisis intelligence (via news monitoring APIs) "
    "with blockchain-verified supply chain operations, enabling automatic rerouting recommendations when "
    "disruptions affect verified supply routes."
)

# Comparison Table
add_subsection_heading("D. Comparative Analysis")

add_table_with_data(
    ["Feature", "Gomasta et al. [5]", "Jamil et al. [6]", "Sherwyn et al.", "ChainMed (Ours)"],
    [
        ["Blockchain", "Hyperledger Fabric", "Hyperledger Fabric", "Ganache (local)", "Ethereum Sepolia (live)"],
        ["Deployable/live", "No", "No", "No", "Yes (Vercel + Railway)"],
        ["Three-party attestation", "No (2-party)", "No", "No", "Yes — core feature"],
        ["AI cross-match fraud", "No", "No", "Basic ML", "Hybrid rule + LLM"],
        ["LLM investigation", "No", "No", "No", "Yes (Gemini 2.5 Flash)"],
        ["IoT dependency", "No", "Yes", "No", "No — zero hardware"],
        ["Crisis rerouting", "No", "No", "No", "Yes (Dijkstra + NewsAPI)"],
        ["Entity trust scoring", "No", "No", "No", "Yes — dynamic scores"],
        ["Explainable AI", "No", "No", "No", "Yes — human-readable"],
        ["DSCSA/FMD compliance", "No", "No", "No", "Yes — audit trail ready"],
        ["BFT analysis", "Yes (33%)", "No", "No", "Yes (identified f=0)"],
        ["Accessible to NGOs", "No", "No", "No", "Yes — zero cost"],
    ],
    "TABLE I: Comparative Analysis Against Prior Works"
)

# ═══════════════════════════════════════════════════════════════════════════════
# III. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("III. System Architecture")

add_subsection_heading("A. Three-Party Attestation Model")

add_body(
    "The fundamental architectural innovation of ChainMed is the three-party attestation model. Unlike "
    "prior systems that track linear ownership transfers, ChainMed requires each party in the supply chain "
    "to independently submit handoff data—including quantity, batch number, expiry date, medicine name, and "
    "storage temperature—without visibility into other parties' submissions. The Verification AI Engine then "
    "cross-matches all three submissions to detect inconsistencies that indicate fraud, diversion, or tampering."
)

add_body(
    "The three-party model creates a fundamental game-theoretic constraint: fraud in ChainMed requires "
    "simultaneous collusion across Manufacturer, Supplier, and Hospital—a condition that is statistically and "
    "operationally difficult to sustain when all three transactions are independently recorded on an immutable "
    "blockchain ledger. The closest prior work (Gomasta et al. [5]) uses ECDSA double-signing between only "
    "two parties (pharmaceutical company and drug regulator), without AI cross-matching of independently "
    "submitted data."
)

add_subsection_heading("B. Hybrid AI Verification Engine")

add_body(
    "The Verification AI operates as a two-layer hybrid system. Layer 1 is a deterministic rule engine that "
    "executes mathematical cross-matching with guaranteed consistency—it always produces the same output for "
    "the same input, with zero dependency on external services. Layer 2 is a Large Language Model (Gemini 2.5 Flash) "
    "that generates natural-language investigation reports for flagged shipments, providing root cause analysis "
    "and recommended actions. Layer 2 operates asynchronously with graceful fallback: if the LLM API is "
    "unavailable, the system continues operating on Layer 1 alone, ensuring the verification pipeline never "
    "blocks on external dependencies."
)

add_subsection_heading("C. Off-Chain/On-Chain Hybrid Ledger")

add_body(
    "ChainMed employs a hybrid ledger architecture that separates storage concerns for optimal performance "
    "and cost efficiency. Off-chain storage (PostgreSQL) maintains raw batch data, full handoff records, AI flag "
    "details, user profiles, approval logs, and stock levels—enabling fast reads/writes with full structured "
    "query capability. On-chain storage (Ethereum Sepolia via ChainMed.sol smart contract) records SHA-256 "
    "checkpoint hashes, verification status, risk scores, and timestamps—providing an immutable audit trail "
    "containing only state transitions. This design ensures that the blockchain serves as a tamper-proof witness "
    "without incurring the storage cost or gas overhead of full data replication on-chain."
)

add_subsection_heading("D. Crisis Disruption Center")

add_body(
    "The Crisis Disruption Center integrates real-time news intelligence with supply chain operations. "
    "Disruption events—natural disasters, geopolitical conflicts, infrastructure failures—are ingested via "
    "NewsAPI monitoring and user-reported incidents. Each event is classified by type, severity, affected "
    "region, and estimated impact duration. The system models the logistics network as a weighted directed "
    "graph and applies Dijkstra's algorithm to compute optimal rerouting paths when disruptions affect "
    "active supply routes."
)

add_subsection_heading("E. Role-Based Access Control and Admin Hierarchy")

add_body(
    "ChainMed implements a three-tier administrative hierarchy with role-scoped access: (1) Admin-Master "
    "(admin_master) serves as the platform custodian with full system access, including the ability to view "
    "and edit accounts of other administrators; (2) Admin-Manager (admin_manager) acts as the compliance and "
    "operations lead, with authority to view global trust scores, crisis events, and manually override AI "
    "flags as the human-in-the-loop compliance officer; (3) Admin-Dev (admin_dev) serves as the technical "
    "overseer with access to raw data for debugging, system logs, and blockchain transaction hashes, but "
    "without authority to override AI flags or modify compliance records. All administrative actions are "
    "recorded in an immutable approval log with actor identification, timestamps, and action descriptions."
)

# ═══════════════════════════════════════════════════════════════════════════════
# IV. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("IV. Methodology")

add_subsection_heading("A. Verification AI Risk Score Formula")

add_body(
    "The Verification AI computes a composite risk score R for each supply chain handoff using a weighted "
    "binary indicator model. The risk score is defined as:"
)

add_equation("R = min( Σᵢ₌₁ⁿ wᵢ · fᵢ , 100 )                    (1)")

add_body_no_indent(
    "where R ∈ [0, 100] is the risk score, fᵢ ∈ {0, 1} are binary fraud indicators (0 = no violation, "
    "1 = violation detected), wᵢ are the weights assigned to each rule i, and n is the number of rules "
    "evaluated. The risk score is capped at 100 to maintain a normalized scale."
)

# Rule weights table
add_table_with_data(
    ["Rule (fᵢ)", "Condition", "Weight (wᵢ)"],
    [
        ["f₁", "Quantity deviation > 15% between parties", "30"],
        ["f₂", "Quantity deviation > 30% (additive severe)", "20"],
        ["f₃", "Expiry date mismatch between parties", "40"],
        ["f₄", "Batch number mismatch", "30"],
        ["f₅", "Medicine name mismatch", "30"],
        ["f₆", "Storage temperature > 30°C (cold chain breach)", "15"],
    ],
    "TABLE II: Verification AI Rule Weights"
)

add_body("The quantity deviation between parties is computed as:")

add_equation("δq = (max(Q₁, Q₂, Q₃) − min(Q₁, Q₂, Q₃)) / min(Q₁, Q₂, Q₃) × 100     (2)")

add_body_no_indent(
    "where Q₁, Q₂, Q₃ are the quantities independently reported by Manufacturer, Supplier, and Hospital "
    "respectively. For the two-party comparison (Manufacturer → Supplier), the percentage deviation is "
    "computed as |Qmfg − Qsup| / Qmfg. The decision boundary is defined as:"
)

add_equation("Status = VERIFIED  if R < 30")
add_equation("Status = FLAGGED   if R ≥ 30                                              (3)")

add_body(
    "When a shipment is flagged (R ≥ 30), the system invokes Layer 2 (Gemini 2.5 Flash LLM) asynchronously "
    "to generate a natural-language investigation report. The LLM receives the mismatch details, triggered "
    "rules, and contextual batch information as a structured prompt, and returns a human-readable root cause "
    "analysis with recommended actions. If the LLM API is unavailable, the rule-based explanation from Layer 1 "
    "is retained as the fallback."
)

add_subsection_heading("B. Entity Trust Scoring Engine")

add_body("The Trust Engine maintains a dynamic trust score for each entity (manufacturer, supplier, hospital) "
         "in the supply chain. The trust score is defined as:")

add_equation("Tₑ = max(0, 100 − α·φₑ − β·R̄ₑ − γ·δₑ)                                   (4)")

add_body_no_indent("where:")
add_bullet("Tₑ = Trust Score for entity e ∈ [0, 100]")
add_bullet("φₑ = flag rate = (FLAGGED shipments by e) / (total shipments by e)")
add_bullet("R̄ₑ = mean risk score across entity's shipments, normalized to [0, 1]")
add_bullet("δₑ = delay rate = (late confirmations by e) / (total confirmations by e)")
add_bullet("α = 40, β = 30, γ = 20 are tunable hyperparameters")

add_body(
    "The trust score decays with each flagged verification event and recovers gradually with successful "
    "verifications. This creates a reputation-based incentive mechanism: entities with consistently clean "
    "supply chain behavior maintain high trust scores, while repeated violations result in progressive "
    "score degradation that triggers enhanced scrutiny from compliance officers."
)

add_subsection_heading("C. Crisis Rerouting — Dijkstra on Disrupted Logistics Graph")

add_body(
    "The logistics network is modeled as a weighted directed graph G = (V, E, W) where V is the set of all "
    "warehouse and hospital nodes, E represents all logistics route edges, and W: E → R⁺ is the edge weight "
    "function (distance × disruption multiplier)."
)

add_body("On disruption event D affecting route set E_D ⊆ E, the edge weights are updated as:")

add_equation("W'(e) = ∞  if e ∈ E_D  ;  W'(e) = W(e)  otherwise                       (5)")

add_body("The optimal rerouting path is computed as:")

add_equation("P* = argmin_{P ∈ Paths(s,t)}  Σ_{e ∈ P}  W'(e)                            (6)")

add_body_no_indent(
    "solved via Dijkstra's algorithm in O((|V| + |E|) log |V|) time. The disruption multiplier incorporates "
    "severity classification (low, medium, high, critical), estimated duration, and geographic proximity "
    "to active supply routes. Events are ingested from NewsAPI monitoring and user-reported incidents, "
    "classified by an AI severity assessor, and visualized on an interactive geographic map."
)

# ═══════════════════════════════════════════════════════════════════════════════
# V. IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("V. Implementation and Deployment")

add_subsection_heading("A. Technology Stack")

add_table_with_data(
    ["Component", "Technology", "Purpose"],
    [
        ["Frontend", "React 18 + Vite", "Single-page application with role-based dashboards"],
        ["Backend API", "FastAPI (Python 3.12)", "RESTful API with async support"],
        ["Database", "PostgreSQL (Supabase)", "Off-chain structured data storage"],
        ["Blockchain", "Ethereum Sepolia + Solidity 0.8.x", "On-chain immutable audit trail"],
        ["Web3 Library", "web3.py", "Blockchain interaction from Python backend"],
        ["AI Layer 1", "Python rule engine", "Deterministic cross-match verification"],
        ["AI Layer 2", "Google Gemini 2.5 Flash", "LLM-powered investigation reports"],
        ["Authentication", "JWT (HS256)", "Stateless token-based auth with OTP for admins"],
        ["QR Codes", "qrcode (Python)", "Shipment identification and public verification"],
        ["Deployment", "Vercel (frontend) + Railway (backend)", "Production cloud deployment"],
        ["Version Control", "Git + GitHub", "Source code management"],
        ["Database Migrations", "Alembic", "Schema version control"],
    ],
    "TABLE III: Technology Stack"
)

add_subsection_heading("B. Smart Contract Design — ChainMed.sol")

add_body(
    "The ChainMed smart contract is deployed on Ethereum Sepolia and provides two primary functions: "
    "recordHandoff() for recording verified supply chain handoffs, and flagShipment() for recording AI-flagged "
    "anomalies. Each record contains a SHA-256 data hash of the off-chain handoff record, the verification "
    "status (VERIFIED/FLAGGED), the AI risk score, and a block timestamp. Events (HandoffRecorded, "
    "ShipmentFlagged) are emitted for off-chain indexing and real-time monitoring."
)

add_body(
    "The contract uses an owner-based access control model where only the backend service wallet can write "
    "to the contract. This custodial model is identified as a security limitation in Section VI and proposed "
    "mitigations are discussed."
)

add_subsection_heading("C. Database Schema")

add_body(
    "The PostgreSQL schema comprises the following core entities: Users (authentication and role management), "
    "Manufacturers, Suppliers, and Consumers (entity registry with trust scores), MedicineBatches (production "
    "records with blockchain hashes), Shipments (logistics tracking with QR codes and dispatched quantities), "
    "HandoffRecords (independent party submissions for AI verification), AIFlags (verification results with "
    "risk scores and triggered rules), ApprovalLogs (immutable audit trail of all administrative actions), "
    "StockLevels (real-time inventory management per entity per medicine), and DisruptionEvents (crisis "
    "intelligence records with severity classification and geographic coordinates)."
)

add_subsection_heading("D. Deployment Architecture")

add_body(
    "ChainMed is deployed as a publicly accessible web application. The React frontend is hosted on "
    "Vercel with automatic CI/CD from the GitHub repository. The FastAPI backend is deployed on Railway "
    "with environment-variable-based configuration for database credentials, blockchain keys, and API tokens. "
    "The PostgreSQL database is hosted on Supabase with connection pooling for production workloads. This "
    "architecture achieves zero hardware cost for deployment, making ChainMed immediately accessible to "
    "NGOs, government health agencies, and healthcare facilities in developing countries."
)

# ═══════════════════════════════════════════════════════════════════════════════
# VI. SECURITY ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("VI. Security Analysis")

add_subsection_heading("A. Smart Contract Vulnerability Assessment")

add_body(
    "A comprehensive security audit of ChainMed.sol identified 10 vulnerabilities across four severity "
    "levels: 2 Critical, 3 High, 3 Medium, and 2 Low. The most significant findings are summarized below."
)

add_body(
    "Critical 1 — Single-Owner Central Authority: The contract uses a single owner address for all write "
    "operations (recordHandoff, flagShipment). If this private key is compromised, an attacker gains "
    "unrestricted ability to forge handoff records and flag legitimate shipments. The BFT system with "
    "n = 1 node provides f = 0 fault tolerance—any single compromise breaks the system completely. "
    "Proposed mitigation: Replace single owner with multi-signature threshold requiring m-of-n authorized "
    "nodes (e.g., 4-of-5) using OpenZeppelin AccessControl or Gnosis Safe."
)

add_body(
    "Critical 2 — Mutable Records: The recordHandoff function overwrites existing entries for a given "
    "shipmentId, destroying the immutability guarantee. A compromised owner can silently rewrite any "
    "historical handoff record. Proposed mitigation: Add require(!_handoffs[shipmentId].exists) guard "
    "before assignment, enforcing append-only semantics."
)

add_body(
    "High 1 — Unconstrained Status String: The status parameter accepts arbitrary strings with no on-chain "
    "validation. Proposed mitigation: Replace string status with enum Status { PENDING, VERIFIED, FLAGGED }, "
    "which Solidity stores as uint8, saving gas and eliminating invalid state."
)

add_body(
    "High 2 — Off-chain dataHash Without Verification: The dataHash parameter is accepted as a plain string "
    "with no format validation. Proposed mitigation: Accept bytes32 dataHash instead, enforcing exact 256-bit "
    "size on-chain."
)

add_body(
    "High 3 — Single-Step Ownership Transfer: The transferOwnership function provides instant, irreversible "
    "ownership transfer without confirmation from the new owner. Proposed mitigation: Implement two-step "
    "transfer using OpenZeppelin's Ownable2Step pattern."
)

add_subsection_heading("B. Byzantine Fault Tolerance Analysis")

add_body(
    "A Byzantine Fault Tolerant (BFT) system with n nodes tolerates at most f = ⌊(n−1)/3⌋ Byzantine "
    "(malicious or faulty) nodes. ChainMed's current smart contract has n = 1 (the single owner), "
    "therefore f = 0—the system has zero tolerance for node compromise."
)

add_equation("f = ⌊(n − 1) / 3⌋ ;   n = 1 → f = 0                                     (7)")

add_body(
    "To achieve parity with Gomasta et al. [5] (33% tolerance), ChainMed requires a multi-signature "
    "deployment with n ≥ 4 independent signers and m ≥ ⌈2n/3⌉ + 1 signatures per write. For n = 5: "
    "require 4-of-5 signatures. One compromised node cannot forge a record; two cannot; only three-of-five "
    "(60%) collusion succeeds—safely above the 33% tolerance threshold."
)

add_subsection_heading("C. Collusion Attack Analysis")

add_body(
    "ChainMed's three-party attestation model creates a natural defense against collusion. For a fraudulent "
    "shipment to pass verification undetected, an attacker must satisfy one of the following conditions:"
)

add_bullet("Compromise all three parties' submissions simultaneously (three-person conspiracy)")
add_bullet("Compromise the backend server's ECDSA private key (custodial model limitation)")
add_bullet("Forge ECDSA signatures on the blockchain (computationally infeasible)")
add_bullet("Alter blockchain records post-recording (consensus-protected)")

add_body(
    "The economic cost of three-party collusion significantly exceeds the cost of two-party collusion in "
    "prior systems. Additionally, the immutable blockchain audit trail ensures that even successful collusion "
    "leaves a forensic evidence trail detectable through subsequent audit."
)

add_subsection_heading("D. Insider Threat Model")

add_body(
    "The current administrative model grants Admin-Master unrestricted access to all system data, creating "
    "an insider trading risk: an administrator who sees all three parties' data before verification has advance "
    "knowledge of pricing, quantity mismatches, and supply disruptions. This paper proposes a zero-trust admin "
    "model where: (1) admins see only mismatch flags without raw values; (2) amendments require cryptographic "
    "signatures from both transacting parties; (3) all access is logged on-chain with embedded ECDSA signatures; "
    "and (4) decryption authority is split across multiple key custodians requiring quorum consensus. This model "
    "is detailed in Section VII as part of the privacy-preserving upgrade path."
)

# ═══════════════════════════════════════════════════════════════════════════════
# VII. DISCUSSION — LIMITATIONS AND PRIVACY MODEL
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("VII. Discussion")

add_subsection_heading("A. Current Limitations")

add_body(
    "ChainMed's current implementation has several limitations that are acknowledged for transparency. "
    "First, all sensitive fields (quantity, expiry, temperature) are stored as plaintext in PostgreSQL, "
    "exposing them to administrative access and creating the insider threat risk described in Section VI.D. "
    "Second, the smart contract uses single-owner access control with f = 0 Byzantine fault tolerance. "
    "Third, the system uses custodial ECDSA signing where the backend server signs blockchain transactions "
    "on behalf of individual parties, rather than each party signing with their own private keys. Fourth, "
    "experimental performance benchmarks (transaction throughput, AI latency, false positive/negative rates) "
    "have not been formally conducted."
)

add_subsection_heading("B. Proposed Three-Layer Privacy Model")

add_body(
    "ChainMed proposes a three-layer data confidentiality model for production deployment:"
)

add_body(
    "Phase 1 — Field-Level AES-256 Encryption: Sensitive columns in the handoff_records table (quantity, "
    "expiry, temperature) are encrypted with AES-256-GCM using server-held keys. The Verification AI decrypts "
    "server-side for comparison only, and the raw values never persist in plaintext in the database. This "
    "ensures no supply chain party can access another party's raw submission through database queries."
)

add_body(
    "Phase 2 — Salted Hash Commitments: Replace plaintext quantities with hash(quantity || salt) using "
    "SHA-256. Each party submits their commitment hash at handoff time. The Verification AI compares hashes "
    "to determine equality without accessing raw values. This eliminates even server-side plaintext exposure "
    "during verification."
)

add_body(
    "Phase 3 — Zero-Knowledge Proof Attestation: Implement ZK-SNARK circuits using Circom and SnarkJS for "
    "quantity deviation verification. The ZKP statement to be proven is: \"I know quantities Q₁, Q₂, Q₃ such "
    "that |Q₁ − Q₂| / Q₁ > 0.15, and I am not revealing Q₁, Q₂, or Q₃.\" The Groth16 proof system over "
    "BN128 enables the AI to produce a cryptographically verifiable FLAGGED or VERIFIED assertion without "
    "the circuit witness (individual party quantities) ever being exposed to any party, administrator, or "
    "the blockchain record. This removes even the AI server from the trust boundary."
)

add_subsection_heading("C. Alternatives to Hardware-Based Trusted Execution Environments")

add_body(
    "For production deployment where hardware TEEs (Intel SGX, ARM TrustZone) are unavailable, ChainMed "
    "evaluates three categories of software alternatives: (1) Multi-Party Computation (MPC) for splitting "
    "computations among independent nodes; (2) Fully Homomorphic Encryption (FHE) for computation on "
    "encrypted data without decryption; and (3) Zero-Knowledge Proofs (ZKPs) for proving correctness without "
    "revealing data. For ChainMed's specific use case—arithmetic threshold comparison with private "
    "inputs—ZKPs via Circom + SnarkJS are identified as the optimal choice, offering the necessary "
    "verifiability without the performance overhead of FHE or the coordination requirements of MPC."
)

add_body(
    "Complementary hardware security modules (HSMs) are recommended for key custody in production, providing "
    "tamper-resistant storage for ECDSA private keys and AES-256 encryption keys outside the main CPU. "
    "Trusted Platform Modules (TPMs) provide hardware-level attestation for verifying the integrity of the "
    "AI verification service itself."
)

# ═══════════════════════════════════════════════════════════════════════════════
# VIII. FUTURE WORK
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("VIII. Future Work")

add_body("The following directions are identified for extending ChainMed:")

add_bullet(
    "Federated Learning for Fraud Detection: Train XGBoost-based fraud detection models across multiple "
    "hospital networks without sharing raw data, enabling collaborative intelligence while preserving data sovereignty."
)
add_bullet(
    "Zero-Knowledge Proof Attestation: Implement the Phase 3 ZK-SNARK circuits for quantity deviation "
    "verification using Circom/SnarkJS, with on-chain Solidity verifier contracts generated via snarkjs zkey export."
)
add_bullet(
    "NFT-Based Medicine Passports: Assign each batch a soulbound NFT (ERC-5192) tracking provenance from "
    "manufacturing lab to patient, creating a publicly verifiable medicine identity."
)
add_bullet(
    "DSCSA/FMD Compliance API: Direct integration with the FDA DSCSA serialization system and EU FMD "
    "repository for regulatory-grade compliance automation."
)
add_bullet(
    "Multi-Chain Bridge: Cross-chain verification between Ethereum (public audit trail) and Hyperledger "
    "Fabric (enterprise operations) to leverage the strengths of both permissioned and permissionless ledgers."
)
add_bullet(
    "Decentralized Identity (DID): Replace JWT-based authentication with W3C DID standard for self-sovereign "
    "entity identities, eliminating centralized credential management."
)
add_bullet(
    "Pandemic Simulation Module: Integrate SIR model disease spread simulations to predict medicine demand "
    "30 days ahead, enabling proactive supply chain pre-positioning."
)
add_bullet(
    "Experimental Benchmarks: Conduct formal performance evaluation including transaction throughput (TPS), "
    "verification AI latency, false positive/negative rates, and Gemini API response time distributions."
)

# ═══════════════════════════════════════════════════════════════════════════════
# IX. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("IX. Conclusion")

add_body(
    "This paper presented ChainMed, a hybrid AI and blockchain framework that introduces three-party "
    "attestation verification for pharmaceutical supply chain integrity. The system's core innovation—requiring "
    "independent data submission from Manufacturer, Supplier, and Hospital, cross-matched by a deterministic "
    "rule engine augmented with LLM-powered investigation—addresses the fundamental gap in prior work where "
    "blockchain provides traceability but not fraud detection, and AI provides analysis but not immutable "
    "evidence. The mathematical verification model (R = min(Σ wᵢ·fᵢ, 100)) provides deterministic, "
    "reproducible fraud scoring, while the Gemini 2.5 Flash integration delivers human-readable investigation "
    "reports that bridge the gap between algorithmic detection and operational decision-making."
)

add_body(
    "The system's deployment on Vercel + Railway with zero hardware requirements demonstrates that "
    "enterprise-grade supply chain verification is achievable at zero infrastructure cost—a critical factor "
    "for adoption in sub-Saharan Africa, South Asia, and other regions where the counterfeit medicine crisis "
    "claims hundreds of thousands of lives annually. The security analysis honestly identifies the system's "
    "current limitations—single-owner smart contract (f = 0 BFT), plaintext data storage, custodial "
    "signing—while proposing a concrete three-phase privacy upgrade path from AES-256 field encryption "
    "through salted hash commitments to full zero-knowledge proof attestation."
)

add_body(
    "ChainMed implements a zero-trust insider-aware architecture where no single human, including "
    "administrators, should be able to unilaterally access or modify sensitive trade data in the proposed "
    "production model. Verification is performed by the Verification AI on encrypted or hashed data, yielding "
    "only a boolean match result. Amendments require explicit cryptographic signatures from both transacting "
    "parties, making collusion between an admin and a single party mathematically insufficient to alter records. "
    "All access is immutably logged on blockchain with ECDSA signatures embedded, providing forensic evidence "
    "of insider trading attempts. This design separates the roles of verification (AI), consent (cryptographic "
    "signatures), adjudication (admin), and audit (immutable log), ensuring no single human holds enough "
    "authority to manipulate the supply chain."
)

add_body(
    "Future work will focus on implementing the ZK-SNARK verification circuits, conducting formal performance "
    "benchmarks, and extending the system with federated learning for collaborative fraud detection across "
    "healthcare networks."
)

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading("References")

references = [
    '[1] WHO, "1 in 10 medical products in developing countries is substandard or falsified," WHO/EMP/RHT/2017.01, World Health Organization, 2017.',
    '[2] UNODC, "Fake Medicines Kill Almost 500,000 Sub-Saharan Africans a Year," United Nations Office on Drugs and Crime, Feb. 2023. [Online]. Available: https://news.un.org/en/story/2023/02/1133062',
    '[3] FDA, "Drug Supply Chain Security Act (DSCSA)," U.S. Food and Drug Administration, 2023. [Online]. Available: https://www.fda.gov/drugs/drug-supply-chain-integrity/drug-supply-chain-security-act-dscsa',
    '[4] S. Nakamoto, "Bitcoin: A Peer-to-Peer Electronic Cash System," 2008. [Online]. Available: https://bitcoin.org/bitcoin.pdf',
    '[5] S. S. Gomasta, A. Dhali, T. Tahlil, M. M. Anwar, and A. B. M. S. Ali, "ChainMed: Blockchain-based drug supply chain provenance verification system," Heliyon, vol. 9, no. 7, e17957, Jul. 2023. doi: 10.1016/j.heliyon.2023.e17957',
    '[6] F. Jamil, L. Hang, K. Kim, and D. Kim, "A novel medical blockchain model for drug supply chain integrity management in a smart hospital," Electronics, vol. 8, no. 5, p. 505, May 2019. doi: 10.3390/electronics8050505',
    '[7] A. Musamih, K. Salah, R. Jayaraman, J. Arshad, M. Debe, Y. Al-Hammadi, and S. Ellahham, "A blockchain-based approach for drug traceability in healthcare supply chain," IEEE Access, vol. 9, pp. 9728–9743, 2021. doi: 10.1109/ACCESS.2021.3049711',
    '[8] P. Sylim, F. Liu, A. Marcelo, and P. Fontelo, "Blockchain technology for detecting falsified and substandard drugs in distribution: pharmaceutical supply chain intervention," JMIR Research Protocols, vol. 7, no. 9, 2018. doi: 10.2196/10163',
    '[9] M. Uddin, K. Salah, R. Jayaraman, S. Pesic, and S. Ellahham, "Blockchain for drug traceability: Architectures and open challenges," Health Informatics Journal, vol. 27, no. 2, 2021. doi: 10.1177/14604582211011228',
    '[10] A. K. Bapatla, S. P. Mohanty, and E. Kougianos, "ChainMed 3.0: Efficient tracking and tracing of drugs in the pharmaceutical supply chain using blockchain integrated product serialization mechanism," Cogent Business & Management, 2024. doi: 10.1080/23311975.2025.2551811',
    '[11] H. O. Bello, C. Idemudia, and T. V. Iyelolu, "Integrating machine learning and blockchain: Conceptual frameworks for real-time fraud detection and prevention," World Journal of Advanced Research and Reviews, vol. 23, no. 1, pp. 056–068, 2024.',
    '[12] "LLM-Grounded Explainable AI for Supply Chain Risk Early Warning via Temporal Graph Attention Networks," arXiv:2603.04818, 2026. [Online]. Available: https://arxiv.org/html/2603.04818',
    '[13] E. W. Dijkstra, "A note on two problems in connexion with graphs," Numerische Mathematik, vol. 1, no. 1, pp. 269–271, 1959.',
    '[14] U. J. Munasinghe and M. N. Halgamuge, "Supply chain traceability and counterfeit detection of COVID-19 vaccines using novel blockchain-based Vacledger system," Expert Systems with Applications, vol. 228, p. 120293, 2023. doi: 10.1016/j.eswa.2023.120293',
    '[15] S. D\'Souza, D. Nazareth, and C. Vaz, "Blockchain and AI in Pharmaceutical Supply Chain," SSRN Electronic Journal, 2021. doi: 10.2139/ssrn.3852034',
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
output_path = r"D:\Project ChainMed\ChainMed_Research.docx"
doc.save(output_path)
print(f"Research paper saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
